import asyncio
import json
import os
import random
import re
import tempfile
from pathlib import Path
from typing import Dict, List, Any, Tuple

from aiogram import Bot, Dispatcher, F
from aiogram.filters import Command
from aiogram.types import Message, CallbackQuery, InlineKeyboardMarkup, InlineKeyboardButton, PollAnswer
from dotenv import load_dotenv

try:
    from docx import Document
except Exception:
    Document = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None


BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
QUESTIONS_FILE = DATA_DIR / "questions.json"
STATS_FILE = DATA_DIR / "user_stats.json"

load_dotenv(BASE_DIR / ".env")

BOT_TOKEN = os.getenv("BOT_TOKEN", "").strip()
ADMIN_IDS = {int(x.strip()) for x in os.getenv("ADMIN_IDS", "").split(",") if x.strip().isdigit()}

if not BOT_TOKEN or BOT_TOKEN == "PASTE_YOUR_BOT_TOKEN_HERE":
    raise RuntimeError("BOT_TOKEN .env faylida yozilmagan. BotFather tokenini kiriting.")

bot = Bot(token=BOT_TOKEN)
dp = Dispatcher()

SESSIONS: Dict[int, Dict[str, Any]] = {}
POLL_MAP: Dict[str, Dict[str, Any]] = {}
ADMIN_UPLOAD: Dict[int, Dict[str, Any]] = {}


def load_questions() -> List[Dict[str, Any]]:
    try:
        with open(QUESTIONS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []


def save_questions(questions: List[Dict[str, Any]]) -> None:
    with open(QUESTIONS_FILE, "w", encoding="utf-8") as f:
        json.dump(questions, f, ensure_ascii=False, indent=2)


def load_stats() -> Dict[str, Any]:
    if not STATS_FILE.exists():
        return {}
    try:
        return json.loads(STATS_FILE.read_text(encoding="utf-8"))
    except Exception:
        return {}


def save_stats(stats: Dict[str, Any]) -> None:
    STATS_FILE.write_text(json.dumps(stats, ensure_ascii=False, indent=2), encoding="utf-8")


def remember_user(user) -> None:
    if not user:
        return

    stats = load_stats()
    uid = str(user.id)

    rec = stats.setdefault(uid, {
        "total_tests": 0,
        "total_questions": 0,
        "total_correct": 0,
        "wrong_ids": []
    })

    full_name = " ".join([x for x in [user.first_name, user.last_name] if x]).strip()
    rec["name"] = full_name or f"User {uid}"
    rec["username"] = user.username or ""

    save_stats(stats)


def clean_text(s: str) -> str:
    s = s.replace("\ufeff", "")
    s = re.sub(r"\r\n?", "\n", s)
    s = re.sub(r"[ \t]+", " ", s)
    return s.strip()


def strip_correct_marks(s: str) -> Tuple[str, bool]:
    original = s
    marks = ["✅", "✔", "[+]", "(+)", "{+}", "#", "*"]
    is_correct = False
    for m in marks:
        if m in s:
            is_correct = True
            s = s.replace(m, "")
    s = re.sub(r"\s+", " ", s).strip(" ;,.\n\t")
    return s.strip(), is_correct


def read_file_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in [".txt", ".csv"]:
        for enc in ["utf-8", "utf-16", "cp1251", "latin-1"]:
            try:
                return path.read_text(encoding=enc)
            except Exception:
                pass
        return path.read_bytes().decode("utf-8", errors="ignore")

    if suffix == ".docx":
        if Document is None:
            raise RuntimeError("python-docx o'rnatilmagan.")
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also read tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append("\n".join(cells))
        return "\n".join(parts)

    if suffix == ".pdf":
        if PdfReader is None:
            raise RuntimeError("pypdf o'rnatilmagan.")
        reader = PdfReader(str(path))
        texts = []
        for page in reader.pages:
            try:
                texts.append(page.extract_text() or "")
            except Exception:
                pass
        return "\n".join(texts)

    raise RuntimeError("Faqat .txt, .docx, .pdf fayllar qabul qilinadi.")


def parse_hemis_format(text: str) -> List[Dict[str, Any]]:
    """Format: question ==== option ==== option ++++ ; first option is usually correct."""
    items = []
    blocks = re.split(r"\n?\+{3,}\n?", text)
    for block in blocks:
        parts = [p.strip() for p in re.split(r"\n?={3,}\n?", block) if p.strip()]
        if len(parts) >= 3:
            question = clean_text(parts[0])
            options = []
            correct_index = 0
            for i, opt in enumerate(parts[1:]):
                opt, mark = strip_correct_marks(clean_text(opt))
                if mark:
                    correct_index = i
                if opt:
                    options.append(opt)
            if question and len(options) >= 2:
                items.append({"question": question, "options": options[:10], "answer_index": correct_index})
    return items


def parse_abcd_format(text: str) -> List[Dict[str, Any]]:
    """
    Supports:
    Question
    A) option
    B) option
    C) option #
    D) option
    Answer: C
    """
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    items = []
    q_lines = []
    opts = []
    correct_letter = None

    opt_re = re.compile(r"^([A-Ja-j])[\)\.\:\-]\s*(.+)$")
    ans_re = re.compile(r"^(?:answer|javob|to.?g.?ri javob|correct)\s*[:\-]\s*([A-Ja-j]|\d+)", re.I)

    def flush():
        nonlocal q_lines, opts, correct_letter
        if q_lines and len(opts) >= 2:
            q = clean_text(" ".join(q_lines))
            options = []
            correct_idx = None
            for idx, (letter, txt, marked) in enumerate(opts):
                txt, mark2 = strip_correct_marks(txt)
                marked = marked or mark2
                if marked:
                    correct_idx = idx
                options.append(txt)
            if correct_letter:
                for idx, (letter, _, _) in enumerate(opts):
                    if letter.upper() == correct_letter.upper():
                        correct_idx = idx
                        break
                if correct_idx is None and correct_letter.isdigit():
                    n = int(correct_letter)
                    if 1 <= n <= len(opts):
                        correct_idx = n - 1
            if correct_idx is None:
                correct_idx = 0
            if q and all(options):
                items.append({"question": q, "options": options[:10], "answer_index": min(correct_idx, 9)})
        q_lines = []
        opts = []
        correct_letter = None

    for ln in lines:
        ans = ans_re.match(ln)
        if ans:
            correct_letter = ans.group(1)
            continue

        m = opt_re.match(ln)
        if m:
            letter = m.group(1).upper()
            txt = m.group(2).strip()
            txt_clean, marked = strip_correct_marks(txt)
            opts.append((letter, txt_clean, marked))
            continue

        # New question starts after options if line is not answer/options
        if opts and q_lines:
            flush()
        q_lines.append(ln)

    flush()
    return items


def parse_numbered_format(text: str) -> List[Dict[str, Any]]:
    """
    Supports:
    1. Question?
       1) option
       2) option
       3) option #
    """
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    normalized = []
    for ln in lines:
        ln = re.sub(r"^\d+[\.\)]\s+(?=.*\?)", "\nQUESTION: ", ln)
        normalized.append(ln)
    return parse_abcd_format("\n".join(normalized))


def parse_questions_from_text(text: str) -> List[Dict[str, Any]]:
    text = clean_text(text)
    all_items = []

    # 1) HEMIS separators
    if "====" in text or "++++" in text:
        all_items.extend(parse_hemis_format(text))

    # 2) A/B/C/D
    all_items.extend(parse_abcd_format(text))

    # Deduplicate by question
    seen = set()
    result = []
    for item in all_items:
        qkey = re.sub(r"\s+", " ", item["question"].lower()).strip()
        if qkey in seen:
            continue
        seen.add(qkey)
        if len(item["question"]) > 300:
            item["question"] = item["question"][:297] + "..."
        item["options"] = [opt[:100] for opt in item["options"] if opt.strip()]
        if len(item["options"]) >= 2:
            item["answer_index"] = max(0, min(int(item.get("answer_index", 0)), len(item["options"]) - 1))
            result.append(item)
    return result


def add_questions_to_db(subject: str, parsed: List[Dict[str, Any]]) -> Tuple[int, int]:
    questions = load_questions()
    max_id = max([int(q.get("id", 0)) for q in questions] or [0])
    existing = {re.sub(r"\s+", " ", q.get("question", "").lower()).strip() for q in questions}

    added = 0
    skipped = 0
    for item in parsed:
        qkey = re.sub(r"\s+", " ", item["question"].lower()).strip()
        if qkey in existing:
            skipped += 1
            continue
        max_id += 1
        questions.append({
            "id": max_id,
            "subject": subject,
            "question": item["question"],
            "options": item["options"],
            "answer_index": item["answer_index"]
        })
        existing.add(qkey)
        added += 1

    save_questions(questions)
    return added, skipped


def get_subjects() -> List[str]:
    return sorted({q.get("subject", "Sun'iy intellekt") for q in load_questions()})


def main_menu() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📚 Fan tanlash", callback_data="menu:subjects")],
        [InlineKeyboardButton(text="📊 Natijalarim", callback_data="menu:stats")],
        [InlineKeyboardButton(text="❌ Xato savollar", callback_data="menu:wrongs")],
    ])


def admin_menu() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="➕ Fayldan fan qo'shish", callback_data="admin:add_subject")],
        [InlineKeyboardButton(text="🔍 Savol qidirish", callback_data="admin:search_question")],
        [InlineKeyboardButton(text="✏️ Savolni ID orqali tahrirlash", callback_data="admin:edit_question_start")],
        [InlineKeyboardButton(text="📥 Backup olish", callback_data="admin:backup")],
        [InlineKeyboardButton(text="📢 Xabar yuborish", callback_data="admin:broadcast_menu")],
        [InlineKeyboardButton(text="🗑 Fanni o'chirish", callback_data="admin:delete_subjects")],
        [InlineKeyboardButton(text="✏️ Fan nomini o'zgartirish", callback_data="admin:rename_subjects")],
        [InlineKeyboardButton(text="🏆 Userlar reytingi", callback_data="admin:ranking")],
        [InlineKeyboardButton(text="👥 Umumiy statistika", callback_data="admin:global_stats")],
        [InlineKeyboardButton(text="📚 Fanlar ro'yxati", callback_data="admin:subjects")],
        [InlineKeyboardButton(text="⬅️ Asosiy menyu", callback_data="menu:home")],
    ])


def subjects_menu() -> InlineKeyboardMarkup:
    rows = [[InlineKeyboardButton(text=f"📘 {sub}", callback_data=f"subject:{sub}")] for sub in get_subjects()]
    rows.append([InlineKeyboardButton(text="⬅️ Orqaga", callback_data="menu:home")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


def admin_subject_action_menu(action: str) -> InlineKeyboardMarkup:
    rows = []
    for sub in get_subjects():
        rows.append([InlineKeyboardButton(text=f"📘 {sub}", callback_data=f"admin:{action}:{sub}")])
    rows.append([InlineKeyboardButton(text="⬅️ Admin menyu", callback_data="admin:home")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


def confirm_delete_subject_menu(subject: str) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Ha, o'chirish", callback_data=f"admin:delete_confirm:{subject}")],
        [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="admin:home")]
    ])


def delete_subject_from_db(subject: str) -> int:
    questions = load_questions()
    before = len(questions)
    questions = [q for q in questions if q.get("subject", "Sun'iy intellekt") != subject]
    save_questions(questions)
    return before - len(questions)


def rename_subject_in_db(old_subject: str, new_subject: str) -> int:
    questions = load_questions()
    changed = 0
    for q in questions:
        if q.get("subject", "Sun'iy intellekt") == old_subject:
            q["subject"] = new_subject
            changed += 1
    save_questions(questions)
    return changed


def format_ranking(limit: int = 20) -> str:
    stats = load_stats()
    if not stats:
        return "🏆 Hali reyting ma'lumoti yo'q."

    rows = []
    for uid, rec in stats.items():
        total_q = int(rec.get("total_questions", 0))
        correct = int(rec.get("total_correct", 0))
        total_tests = int(rec.get("total_tests", 0))
        percent = round(correct / total_q * 100, 1) if total_q else 0

        name = rec.get("name") or f"User {uid}"
        username = rec.get("username") or ""
        display = name
        if username:
            display += f" (@{username})"

        rows.append({
            "user_id": uid,
            "display": display,
            "total_tests": total_tests,
            "total_questions": total_q,
            "correct": correct,
            "percent": percent
        })

    # Avval foiz, keyin to'g'ri javob, keyin test soni bo'yicha tartiblaydi
    rows.sort(key=lambda x: (x["percent"], x["correct"], x["total_tests"]), reverse=True)

    medals = ["🥇", "🥈", "🥉"]
    text = "🏆 Userlar reytingi\n\n"
    for i, row in enumerate(rows[:limit], start=1):
        medal = medals[i-1] if i <= 3 else f"{i}."
        text += (
            f"{medal} {row['display']}\n"
            f"   📈 Foiz: {row['percent']}%\n"
            f"   ✅ To'g'ri: {row['correct']}/{row['total_questions']}\n"
            f"   📝 Testlar: {row['total_tests']}\n"
            f"   🆔 ID: {row['user_id']}\n\n"
        )
    return text.strip()


def format_global_stats() -> str:
    stats = load_stats()
    questions = load_questions()

    subjects = {}
    for q in questions:
        s = q.get("subject", "Noma'lum")
        subjects[s] = subjects.get(s, 0) + 1

    total_users = len(stats)
    active_users = sum(1 for rec in stats.values() if int(rec.get("total_questions", 0)) > 0)
    total_tests = sum(int(rec.get("total_tests", 0)) for rec in stats.values())
    total_answered = sum(int(rec.get("total_questions", 0)) for rec in stats.values())
    total_correct = sum(int(rec.get("total_correct", 0)) for rec in stats.values())
    percent = round(total_correct / total_answered * 100, 1) if total_answered else 0

    text = (
        "👥 Umumiy statistika\n\n"
        f"👤 Jami userlar: {total_users}\n"
        f"🔥 Test ishlagan userlar: {active_users}\n"
        f"📝 Jami testlar: {total_tests}\n"
        f"📌 Jami javob berilgan savollar: {total_answered}\n"
        f"✅ Jami to'g'ri javoblar: {total_correct}\n"
        f"📈 Umumiy foiz: {percent}%\n\n"
        f"📚 Fanlar soni: {len(subjects)}\n"
        f"🧾 Bazadagi savollar: {len(questions)}\n\n"
        "📚 Fanlar bo'yicha savollar:\n"
    )

    if subjects:
        for sub, count in sorted(subjects.items()):
            text += f"• {sub}: {count} ta\n"
    else:
        text += "Fanlar yo'q.\n"
    return text.strip()


def get_question_by_id(qid: int):
    for q in load_questions():
        if int(q.get("id", -1)) == qid:
            return q
    return None


def format_question_for_admin(q: Dict[str, Any]) -> str:
    letters = "ABCDEFGHIJ"
    options = q.get("options", [])
    ans = int(q.get("answer_index", 0))
    text = (
        f"🆔 ID: {q.get('id')}\n"
        f"📚 Fan: {q.get('subject', 'Nomaʼlum')}\n\n"
        f"❓ Savol:\n{q.get('question', '')}\n\n"
        "Variantlar:\n"
    )
    for i, opt in enumerate(options):
        mark = " ✅" if i == ans else ""
        letter = letters[i] if i < len(letters) else str(i + 1)
        text += f"{letter}) {opt}{mark}\n"
    if options and 0 <= ans < len(options):
        text += f"\n✅ To'g'ri javob: {letters[ans] if ans < len(letters) else ans + 1}"
    else:
        text += "\n⚠️ To'g'ri javob noto'g'ri ko'rsatilgan"
    return text


def question_edit_menu(qid: int) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✏️ Savol matnini o'zgartirish", callback_data=f"admin:qedit_text:{qid}")],
        [InlineKeyboardButton(text="🔤 Variantlarni o'zgartirish", callback_data=f"admin:qedit_options:{qid}")],
        [InlineKeyboardButton(text="✅ To'g'ri javobni o'zgartirish", callback_data=f"admin:qedit_answer:{qid}")],
        [InlineKeyboardButton(text="🗑 Savolni o'chirish", callback_data=f"admin:qdelete_ask:{qid}")],
        [InlineKeyboardButton(text="⬅️ Admin menyu", callback_data="admin:home")]
    ])


def answer_select_menu(qid: int, options_count: int) -> InlineKeyboardMarkup:
    letters = "ABCDEFGHIJ"
    rows = []
    for i in range(options_count):
        letter = letters[i] if i < len(letters) else str(i + 1)
        rows.append([InlineKeyboardButton(text=f"{letter}) {i+1}-variant", callback_data=f"admin:qset_answer:{qid}:{i}")])
    rows.append([InlineKeyboardButton(text="⬅️ Orqaga", callback_data=f"admin:qshow:{qid}")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


def confirm_delete_question_menu(qid: int) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Ha, savolni o'chirish", callback_data=f"admin:qdelete_confirm:{qid}")],
        [InlineKeyboardButton(text="❌ Bekor qilish", callback_data=f"admin:qshow:{qid}")]
    ])


def update_question_text(qid: int, new_text: str) -> bool:
    questions = load_questions()
    for q in questions:
        if int(q.get("id", -1)) == qid:
            q["question"] = new_text
            save_questions(questions)
            return True
    return False


def update_question_options(qid: int, options: List[str]) -> bool:
    questions = load_questions()
    for q in questions:
        if int(q.get("id", -1)) == qid:
            old_answer = int(q.get("answer_index", 0))
            q["options"] = options
            q["answer_index"] = min(old_answer, len(options) - 1)
            save_questions(questions)
            return True
    return False


def update_question_answer(qid: int, answer_index: int) -> bool:
    questions = load_questions()
    for q in questions:
        if int(q.get("id", -1)) == qid:
            if 0 <= answer_index < len(q.get("options", [])):
                q["answer_index"] = answer_index
                save_questions(questions)
                return True
    return False


def delete_question_by_id(qid: int) -> bool:
    questions = load_questions()
    before = len(questions)
    questions = [q for q in questions if int(q.get("id", -1)) != qid]
    if len(questions) != before:
        save_questions(questions)
        return True
    return False


def search_questions(keyword: str, limit: int = 10) -> List[Dict[str, Any]]:
    keyword = keyword.lower().strip()
    if not keyword:
        return []
    result = []
    for q in load_questions():
        text = " ".join([
            str(q.get("question", "")),
            str(q.get("subject", "")),
            " ".join(q.get("options", []))
        ]).lower()
        if keyword in text:
            result.append(q)
            if len(result) >= limit:
                break
    return result


def search_results_menu(results: List[Dict[str, Any]]) -> InlineKeyboardMarkup:
    rows = []
    for q in results:
        question = q.get("question", "")
        title = question[:45] + ("..." if len(question) > 45 else "")
        rows.append([InlineKeyboardButton(text=f"ID {q.get('id')}: {title}", callback_data=f"admin:qshow:{q.get('id')}")])
    rows.append([InlineKeyboardButton(text="⬅️ Admin menyu", callback_data="admin:home")])
    return InlineKeyboardMarkup(inline_keyboard=rows)



def broadcast_menu() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📢 Hamma userlarga", callback_data="admin:broadcast_all")],
        [InlineKeyboardButton(text="🔥 Aktiv userlarga", callback_data="admin:broadcast_active")],
        [InlineKeyboardButton(text="👤 Bitta user ID ga", callback_data="admin:broadcast_one")],
        [InlineKeyboardButton(text="⬅️ Admin menyu", callback_data="admin:home")]
    ])


def broadcast_confirm_menu() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="✅ Yuborish", callback_data="admin:broadcast_confirm")],
        [InlineKeyboardButton(text="❌ Bekor qilish", callback_data="admin:broadcast_cancel")]
    ])


def get_broadcast_targets(mode: str, target_user_id: str = "") -> List[int]:
    stats = load_stats()
    targets = []

    if mode == "all":
        for uid in stats.keys():
            if str(uid).isdigit():
                targets.append(int(uid))

    elif mode == "active":
        for uid, rec in stats.items():
            if str(uid).isdigit() and int(rec.get("total_questions", 0)) > 0:
                targets.append(int(uid))

    elif mode == "one":
        if str(target_user_id).isdigit():
            targets.append(int(target_user_id))

    # adminlarga alohida yuborish funksiyasi yo'q; lekin admin user sifatida statistikada bo'lsa, all/active ichida chiqishi mumkin.
    # Dublikatlarni olib tashlaymiz
    return sorted(set(targets))


def broadcast_mode_title(mode: str) -> str:
    if mode == "all":
        return "📢 Hamma userlar"
    if mode == "active":
        return "🔥 Aktiv userlar"
    if mode == "one":
        return "👤 Bitta user"
    return "📢 Xabar"


def settings_menu(subject: str) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="⏱ 30 soniya", callback_data=f"time:{subject}:30"),
         InlineKeyboardButton(text="⏱ 1 daqiqa", callback_data=f"time:{subject}:60")],
        [InlineKeyboardButton(text="⬅️ Fanlarga qaytish", callback_data="menu:subjects")]
    ])


def blocks_menu(subject: str, seconds: int) -> InlineKeyboardMarkup:
    qs = [q for q in load_questions() if q.get("subject", "Sun'iy intellekt") == subject]
    rows = []
    for start in range(0, len(qs), 30):
        end = min(start + 30, len(qs))
        rows.append([InlineKeyboardButton(text=f"📝 {start+1}-{end} savollar", callback_data=f"start:{subject}:{seconds}:{start}:{end}")])
    rows.append([InlineKeyboardButton(text="🎲 Random 30 ta", callback_data=f"random:{subject}:{seconds}")])
    rows.append([InlineKeyboardButton(text="⬅️ Vaqt tanlash", callback_data=f"subject:{subject}")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


def after_result_menu(has_wrongs: bool) -> InlineKeyboardMarkup:
    rows = [[InlineKeyboardButton(text="🔁 Yana test ishlash", callback_data="menu:subjects")],
            [InlineKeyboardButton(text="📊 Natijalarim", callback_data="menu:stats")]]
    if has_wrongs:
        rows.insert(0, [InlineKeyboardButton(text="❌ Xato savollarni qayta ishlash", callback_data="menu:wrongs")])
    return InlineKeyboardMarkup(inline_keyboard=rows)


def normalize_for_poll(q: Dict[str, Any]) -> Dict[str, Any]:
    options = list(q["options"])
    original_correct = int(q.get("answer_index", 0))
    pairs = [(opt, idx == original_correct) for idx, opt in enumerate(options)]
    random.shuffle(pairs)
    q2 = dict(q)
    q2["poll_options"] = [p[0] for p in pairs]
    q2["poll_correct_option_id"] = next(i for i, p in enumerate(pairs) if p[1])
    return q2


def update_user_stats(user_id: int, total: int, correct: int, wrong_ids: List[int]) -> None:
    stats = load_stats()
    uid = str(user_id)
    rec = stats.setdefault(uid, {"total_tests": 0, "total_questions": 0, "total_correct": 0, "wrong_ids": []})
    rec["total_tests"] += 1
    rec["total_questions"] += total
    rec["total_correct"] += correct
    saved = set(rec.get("wrong_ids", []))
    saved.update(wrong_ids)
    session = SESSIONS.get(user_id)
    if session:
        saved -= set(session.get("correct_ids", []))
    rec["wrong_ids"] = sorted(saved)
    save_stats(stats)


async def send_next_question(user_id: int) -> None:
    session = SESSIONS.get(user_id)
    if not session:
        return
    index = session["index"]
    questions = session["questions"]

    if index >= len(questions):
        total = len(questions)
        correct = session["correct"]
        wrong_ids = session["wrong_ids"]
        percent = round(correct / total * 100, 1) if total else 0
        update_user_stats(user_id, total, correct, wrong_ids)
        SESSIONS.pop(user_id, None)
        await bot.send_message(
            user_id,
            f"✅ Test tugadi!\n\n📚 Fan: {session['subject']}\n📌 Savollar: {total} ta\n✅ To'g'ri: {correct} ta\n❌ Xato: {len(wrong_ids)} ta\n📈 Foiz: {percent}%",
            reply_markup=after_result_menu(bool(wrong_ids))
        )
        return

    q = normalize_for_poll(questions[index])
    session["current_question"] = q

    title = f"{index+1}/{len(questions)}. {q['question']}"
    title = title[:300]
    opts = [str(x).strip()[:100] for x in q["poll_options"]]

    msg = await bot.send_poll(
        chat_id=user_id,
        question=title,
        options=opts,
        type="quiz",
        correct_option_id=q["poll_correct_option_id"],
        is_anonymous=False,
        open_period=session["seconds"],
        explanation=f"To'g'ri javob: {opts[q['poll_correct_option_id']]}"[:200],
    )

    POLL_MAP[msg.poll.id] = {"user_id": user_id, "question": q}
    asyncio.create_task(auto_next_after_timeout(user_id, msg.poll.id, session["seconds"] + 2))


async def auto_next_after_timeout(user_id: int, poll_id: str, delay: int) -> None:
    await asyncio.sleep(delay)
    session = SESSIONS.get(user_id)
    data = POLL_MAP.get(poll_id)
    if not session or not data:
        return
    current = session.get("current_question")
    if current and current.get("id") == data["question"].get("id"):
        qid = int(current["id"])
        if qid not in session["answered_ids"]:
            session["wrong_ids"].append(qid)
            session["answered_ids"].add(qid)
            session["index"] += 1
            await bot.send_message(user_id, "⏰ Vaqt tugadi. Keyingi savol.")
            await send_next_question(user_id)


@dp.message(Command("start"))
async def cmd_start(message: Message):
    remember_user(message.from_user)
    await message.answer("Assalomu alaykum!\n\nBoshlash uchun fan tanlang.", reply_markup=main_menu())


@dp.message(Command("admin"))
async def cmd_admin(message: Message):
    if message.from_user.id not in ADMIN_IDS:
        await message.answer("Bu menyu faqat admin uchun.")
        return
    await message.answer("Admin menyu:", reply_markup=admin_menu())


@dp.callback_query(F.data == "admin:add_subject")
async def cb_admin_add(call: CallbackQuery):
    if call.from_user.id not in ADMIN_IDS:
        await call.answer("Admin emas.", show_alert=True)
        return
    ADMIN_UPLOAD[call.from_user.id] = {"step": "subject"}
    await call.message.answer("Yangi fan nomini yozing. Masalan: Kiberxavfsizlik")
    await call.answer()


@dp.callback_query(F.data == "admin:subjects")
async def cb_admin_subjects(call: CallbackQuery):
    if call.from_user.id not in ADMIN_IDS:
        await call.answer("Admin emas.", show_alert=True)
        return
    qs = load_questions()
    counts = {}
    for q in qs:
        counts[q.get("subject", "Noma'lum")] = counts.get(q.get("subject", "Noma'lum"), 0) + 1
    text = "📚 Fanlar:\n\n" + "\n".join([f"• {k}: {v} ta savol" for k, v in counts.items()])
    await call.message.answer(text or "Fanlar topilmadi.")
    await call.answer()



@dp.callback_query(F.data == "admin:home")
async def cb_admin_home(call: CallbackQuery):
    if call.from_user.id not in ADMIN_IDS:
        await call.answer("Admin emas.", show_alert=True)
        return
    await call.message.edit_text("Admin menyu:", reply_markup=admin_menu())
    await call.answer()


@dp.callback_query(F.data == "admin:delete_subjects")
async def cb_admin_delete_subjects(call: CallbackQuery):
    if call.from_user.id not in ADMIN_IDS:
        await call.answer("Admin emas.", show_alert=True)
        return
    await call.message.edit_text("Qaysi fanni o'chirmoqchisiz?", reply_markup=admin_subject_action_menu("delete_select"))
    await call.answer()


@dp.callback_query(F.data.startswith("admin:delete_select:"))
async def cb_admin_delete_select(call: CallbackQuery):
    if call.from_user.id not in ADMIN_IDS:
        await call.answer("Admin emas.", show_alert=True)
        return
    subject = call.data.split(":", 2)[2]
    count = sum(1 for q in load_questions() if q.get("subject", "Sun'iy intellekt") == subject)
    await call.message.edit_text(
        f"⚠️ Diqqat!\n\n"
        f"Fan: {subject}\n"
        f"Savollar soni: {count} ta\n\n"
        f"Haqiqatan ham shu fanni bazadan butunlay o'chirasizmi?",
        reply_markup=confirm_delete_subject_menu(subject)
    )
    await call.answer()


@dp.callback_query(F.data.startswith("admin:delete_confirm:"))
async def cb_admin_delete_confirm(call: CallbackQuery):
    if call.from_user.id not in ADMIN_IDS:
        await call.answer("Admin emas.", show_alert=True)
        return
    subject = call.data.split(":", 2)[2]
    deleted = delete_subject_from_db(subject)
    await call.message.edit_text(
        f"✅ Fan o'chirildi.\n\n"
        f"Fan: {subject}\n"
        f"O'chirilgan savollar: {deleted} ta",
        reply_markup=admin_menu()
    )
    await call.answer()


@dp.callback_query(F.data == "admin:rename_subjects")
async def cb_admin_rename_subjects(call: CallbackQuery):
    if call.from_user.id not in ADMIN_IDS:
        await call.answer("Admin emas.", show_alert=True)
        return
    await call.message.edit_text("Qaysi fan nomini o'zgartirmoqchisiz?", reply_markup=admin_subject_action_menu("rename_select"))
    await call.answer()


@dp.callback_query(F.data.startswith("admin:rename_select:"))
async def cb_admin_rename_select(call: CallbackQuery):
    if call.from_user.id not in ADMIN_IDS:
        await call.answer("Admin emas.", show_alert=True)
        return
    subject = call.data.split(":", 2)[2]
    ADMIN_UPLOAD[call.from_user.id] = {"step": "rename_subject", "old_subject": subject}
    await call.message.answer(
        f"Eski fan nomi: {subject}\n\n"
        "Yangi fan nomini yozing:"
    )
    await call.answer()


@dp.callback_query(F.data == "admin:ranking")
async def cb_admin_ranking(call: CallbackQuery):
    if call.from_user.id not in ADMIN_IDS:
        await call.answer("Admin emas.", show_alert=True)
        return
    await call.message.answer(format_ranking(), reply_markup=admin_menu())
    await call.answer()


@dp.callback_query(F.data == "admin:global_stats")
async def cb_admin_global_stats(call: CallbackQuery):
    if call.from_user.id not in ADMIN_IDS:
        await call.answer("Admin emas.", show_alert=True)
        return
    await call.message.answer(format_global_stats(), reply_markup=admin_menu())
    await call.answer()


@dp.message(F.document)
async def upload_file(message: Message):
    if message.from_user.id not in ADMIN_IDS:
        return
    state = ADMIN_UPLOAD.get(message.from_user.id)
    if not state or state.get("step") != "file":
        await message.answer("Fayl qo'shish uchun avval /admin → Fayldan fan qo'shish ni bosing.")
        return

    doc = message.document
    filename = doc.file_name or "uploaded.txt"
    suffix = Path(filename).suffix.lower()
    if suffix not in [".txt", ".docx", ".pdf", ".csv"]:
        await message.answer("Faqat .txt, .docx, .pdf yoki .csv fayl yuboring.")
        return

    with tempfile.TemporaryDirectory() as td:
        path = Path(td) / filename
        await bot.download(doc, destination=path)
        try:
            text = read_file_text(path)
            parsed = parse_questions_from_text(text)
        except Exception as e:
            await message.answer(f"Faylni o'qishda xato: {e}")
            return

    if not parsed:
        await message.answer(
            "Savollar topilmadi.\n\n"
            "Eng yaxshi format:\n"
            "Savol?\n====\nTo'g'ri javob\n====\nXato javob\n====\nXato javob\n++++\n\n"
            "Yoki:\nSavol?\nA) javob #\nB) javob\nC) javob"
        )
        return

    subject = state["subject"]
    added, skipped = add_questions_to_db(subject, parsed)
    ADMIN_UPLOAD.pop(message.from_user.id, None)

    await message.answer(
        f"✅ Fan qo'shildi: {subject}\n\n"
        f"📄 Fayldan topildi: {len(parsed)} ta savol\n"
        f"➕ Bazaga qo'shildi: {added} ta\n"
        f"♻️ Takror savol o'tkazib yuborildi: {skipped} ta\n\n"
        "Endi /start orqali fanlar ro'yxatida chiqadi.",
        reply_markup=admin_menu()
    )


@dp.message(Command("edit"))
async def cmd_edit(message: Message):
    if message.from_user.id not in ADMIN_IDS:
        await message.answer("Bu buyruq faqat admin uchun.")
        return
    parts = message.text.split()
    if len(parts) != 3 or not parts[1].isdigit() or not parts[2].isdigit():
        await message.answer("Format: /edit SAVOL_ID TOGRI_VARIANT_RAQAMI\nMasalan: /edit 29 3")
        return
    qid, variant_no = int(parts[1]), int(parts[2])
    qs = load_questions()
    for q in qs:
        if int(q.get("id", -1)) == qid:
            if not 1 <= variant_no <= len(q["options"]):
                await message.answer(f"Variant 1 dan {len(q['options'])} gacha bo'lishi kerak.")
                return
            q["answer_index"] = variant_no - 1
            save_questions(qs)
            await message.answer(f"✅ ID {qid}: to'g'ri javob {variant_no}-variant qilindi.")
            return
    await message.answer("Bunday ID topilmadi.")


@dp.message(Command("stop"))
async def cmd_stop(message: Message):
    SESSIONS.pop(message.from_user.id, None)
    ADMIN_UPLOAD.pop(message.from_user.id, None)
    await message.answer("Jarayon to'xtatildi.", reply_markup=main_menu())


@dp.message()
async def text_handler(message: Message):
    if message.from_user.id in ADMIN_IDS:
        state = ADMIN_UPLOAD.get(message.from_user.id)

        if state and state.get("step") == "subject":
            subject = message.text.strip()
            if len(subject) < 2:
                await message.answer("Fan nomi juda qisqa. Qayta yozing.")
                return
            state["subject"] = subject
            state["step"] = "file"
            await message.answer(
                f"Fan nomi: {subject}\n\nEndi test faylini yuboring: .txt, .docx yoki .pdf"
            )
            return

        if state and state.get("step") == "rename_subject":
            new_subject = message.text.strip()
            if len(new_subject) < 2:
                await message.answer("Yangi fan nomi juda qisqa. Qayta yozing.")
                return
            old_subject = state["old_subject"]
            changed = rename_subject_in_db(old_subject, new_subject)
            ADMIN_UPLOAD.pop(message.from_user.id, None)
            await message.answer(
                f"✅ Fan nomi o'zgartirildi.\n\n"
                f"Eski nom: {old_subject}\n"
                f"Yangi nom: {new_subject}\n"
                f"O'zgargan savollar: {changed} ta",
                reply_markup=admin_menu()
            )
            return

        if state and state.get("step") == "search_question":
            keyword = message.text.strip()
            results = search_questions(keyword)
            ADMIN_UPLOAD.pop(message.from_user.id, None)
            if not results:
                await message.answer("Hech narsa topilmadi.", reply_markup=admin_menu())
                return
            await message.answer(
                f"🔍 Topildi: {len(results)} ta natija.\nKerakli savolni tanlang:",
                reply_markup=search_results_menu(results)
            )
            return

        if state and state.get("step") == "edit_question_id":
            text = message.text.strip()
            if not text.isdigit():
                await message.answer("Faqat ID raqam yuboring. Masalan: 25")
                return
            qid = int(text)
            q = get_question_by_id(qid)
            ADMIN_UPLOAD.pop(message.from_user.id, None)
            if not q:
                await message.answer("Bunday ID topilmadi.", reply_markup=admin_menu())
                return
            await message.answer(format_question_for_admin(q), reply_markup=question_edit_menu(qid))
            return

        if state and state.get("step") == "edit_question_text":
            qid = int(state["qid"])
            new_text = message.text.strip()
            if len(new_text) < 3:
                await message.answer("Savol matni juda qisqa. Qayta yuboring.")
                return
            ok = update_question_text(qid, new_text)
            ADMIN_UPLOAD.pop(message.from_user.id, None)
            q = get_question_by_id(qid)
            if ok and q:
                await message.answer("✅ Savol matni yangilandi.\n\n" + format_question_for_admin(q), reply_markup=question_edit_menu(qid))
            else:
                await message.answer("Savol topilmadi.", reply_markup=admin_menu())
            return

        if state and state.get("step") == "edit_question_options":
            qid = int(state["qid"])
            lines = [x.strip() for x in message.text.splitlines() if x.strip()]
            options = []
            for line in lines:
                line = re.sub(r"^[A-Ja-j][\)\.\:\-]\s*", "", line).strip()
                if line:
                    options.append(line[:100])
            if len(options) < 2:
                await message.answer("Kamida 2 ta variant yuboring.")
                return
            if len(options) > 10:
                options = options[:10]
            ok = update_question_options(qid, options)
            ADMIN_UPLOAD.pop(message.from_user.id, None)
            q = get_question_by_id(qid)
            if ok and q:
                await message.answer("✅ Variantlar yangilandi.\n\n" + format_question_for_admin(q), reply_markup=question_edit_menu(qid))
            else:
                await message.answer("Savol topilmadi.", reply_markup=admin_menu())
            return


        if state and state.get("step") == "broadcast_one_id":
            target_id = message.text.strip()
            if not target_id.isdigit():
                await message.answer("User ID faqat raqam bo'lishi kerak. Qayta yuboring.")
                return
            state["target_user_id"] = target_id
            state["mode"] = "one"
            state["step"] = "broadcast_text"
            await message.answer(
                f"👤 User ID: {target_id}\n\n"
                "Endi yuboriladigan xabar matnini yozing:"
            )
            return

        if state and state.get("step") == "broadcast_text":
            text = message.text.strip()
            if len(text) < 2:
                await message.answer("Xabar juda qisqa. Qayta yozing.")
                return

            state["text"] = text
            state["step"] = "broadcast_confirm"

            mode = state.get("mode")
            target_user_id = state.get("target_user_id", "")
            targets = get_broadcast_targets(mode, target_user_id)

            preview = (
                f"📢 Xabar preview\n\n"
                f"📌 Rejim: {broadcast_mode_title(mode)}\n"
                f"👥 Qabul qiluvchilar: {len(targets)} ta\n\n"
                f"Xabar:\n{text}\n\n"
                "Yuborishni tasdiqlaysizmi?"
            )
            await message.answer(preview, reply_markup=broadcast_confirm_menu())
            return


@dp.callback_query(F.data == "menu:home")
async def cb_home(call: CallbackQuery):
    await call.message.edit_text("Asosiy menyu:", reply_markup=main_menu())
    await call.answer()


@dp.callback_query(F.data == "menu:subjects")
async def cb_subjects(call: CallbackQuery):
    await call.message.edit_text("Fan nomini tanlang:", reply_markup=subjects_menu())
    await call.answer()


@dp.callback_query(F.data.startswith("subject:"))
async def cb_subject(call: CallbackQuery):
    subject = call.data.split(":", 1)[1]
    await call.message.edit_text(f"📘 Fan: {subject}\n\nVaqtni tanlang:", reply_markup=settings_menu(subject))
    await call.answer()


@dp.callback_query(F.data.startswith("time:"))
async def cb_time(call: CallbackQuery):
    _, subject, seconds = call.data.split(":", 2)
    await call.message.edit_text(f"📘 Fan: {subject}\n⏱ Vaqt: {seconds} soniya\n\nTest bo'limini tanlang:", reply_markup=blocks_menu(subject, int(seconds)))
    await call.answer()


@dp.callback_query(F.data.startswith("start:"))
async def cb_start_block(call: CallbackQuery):
    _, subject, seconds, start, end = call.data.split(":", 4)
    seconds, start, end = int(seconds), int(start), int(end)
    qs = [q for q in load_questions() if q.get("subject", "Sun'iy intellekt") == subject]
    selected = qs[start:end]
    if not selected:
        await call.answer("Savollar topilmadi.", show_alert=True)
        return
    SESSIONS[call.from_user.id] = {"subject": subject, "seconds": seconds, "questions": selected, "index": 0, "correct": 0, "wrong_ids": [], "correct_ids": [], "answered_ids": set(), "current_question": None}
    await call.message.answer(f"🚀 Test boshlandi!\n📘 Fan: {subject}\n📌 Savollar: {start+1}-{end}\n⏱ Har savol: {seconds} soniya")
    await call.answer()
    await send_next_question(call.from_user.id)


@dp.callback_query(F.data.startswith("random:"))
async def cb_random(call: CallbackQuery):
    _, subject, seconds = call.data.split(":", 2)
    seconds = int(seconds)
    qs = [q for q in load_questions() if q.get("subject", "Sun'iy intellekt") == subject]
    selected = random.sample(qs, min(30, len(qs)))
    SESSIONS[call.from_user.id] = {"subject": subject, "seconds": seconds, "questions": selected, "index": 0, "correct": 0, "wrong_ids": [], "correct_ids": [], "answered_ids": set(), "current_question": None}
    await call.message.answer(f"🎲 Random test boshlandi!\n📘 Fan: {subject}\n📌 Savollar: {len(selected)} ta\n⏱ Har savol: {seconds} soniya")
    await call.answer()
    await send_next_question(call.from_user.id)


@dp.callback_query(F.data == "menu:stats")
async def cb_stats(call: CallbackQuery):
    rec = load_stats().get(str(call.from_user.id), {})
    tq = rec.get("total_questions", 0)
    tc = rec.get("total_correct", 0)
    percent = round(tc / tq * 100, 1) if tq else 0
    await call.message.answer(
        f"📊 Natijalaringiz:\n\n📝 Testlar soni: {rec.get('total_tests', 0)}\n📌 Jami savollar: {tq}\n✅ To'g'ri javoblar: {tc}\n❌ Saqlangan xato savollar: {len(rec.get('wrong_ids', []))}\n📈 Umumiy foiz: {percent}%",
        reply_markup=main_menu()
    )
    await call.answer()


@dp.callback_query(F.data == "menu:wrongs")
async def cb_wrongs(call: CallbackQuery):
    wrong_ids = load_stats().get(str(call.from_user.id), {}).get("wrong_ids", [])
    if not wrong_ids:
        await call.message.answer("Sizda hozircha xato savollar yo'q.", reply_markup=main_menu())
        await call.answer()
        return
    qs = [q for q in load_questions() if int(q.get("id", -1)) in set(wrong_ids)]
    SESSIONS[call.from_user.id] = {"subject": "Xato savollar", "seconds": 30, "questions": qs, "index": 0, "correct": 0, "wrong_ids": [], "correct_ids": [], "answered_ids": set(), "current_question": None}
    await call.message.answer(f"❌ Xato savollar testi boshlandi. Jami: {len(qs)} ta")
    await call.answer()
    await send_next_question(call.from_user.id)


@dp.poll_answer()
async def poll_answer_handler(answer: PollAnswer):
    remember_user(answer.user)
    data = POLL_MAP.get(answer.poll_id)
    if not data:
        return
    user_id = answer.user.id
    if data["user_id"] != user_id:
        return
    session = SESSIONS.get(user_id)
    if not session:
        return
    q = data["question"]
    qid = int(q["id"])
    if qid in session["answered_ids"]:
        return
    selected = answer.option_ids[0] if answer.option_ids else -1
    if selected == int(q["poll_correct_option_id"]):
        session["correct"] += 1
        session["correct_ids"].append(qid)
    else:
        session["wrong_ids"].append(qid)
    session["answered_ids"].add(qid)
    session["index"] += 1
    await asyncio.sleep(1)
    await send_next_question(user_id)


async def main():
    print("Bot started with admin file upload parser...")
    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())
